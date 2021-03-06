#!/usr/bin/python
# -*- coding: utf-8 -*-

#!/usr/bin/python
# -*- coding: utf-8 -*-
import datetime
from datetime import date, timedelta
import time
import StringIO
import xlsxwriter

from django.db.models import Avg, Sum, Max, Min, Count
from django.conf import settings
from django.conf.urls.static import static
from django.http import HttpResponse,HttpResponseRedirect, Http404
from django.utils.translation import ugettext, ugettext_lazy as _

from docx import Document
from docx.shared import Inches, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH

from personel.models import *
from personel.helpScripts import td2DayHourMin

from .utils import *
#from .utils import get_temperatures, get_wind_speed, get_str_days, get_random_colors, precip_prob_sum, get_percentage, legendcolors = get_random_colors(10)



#########################################
# DOCX VIEWS
#########################################
#url(r'^doc/test/$', staff_or_403(docTest), name='docTest'),
def docTest(request, LessonID=None):
    
    #Lesson as filter
    LessonID = request.GET.get('LessonID', None)
    if LessonID:
        lesson = Lesson.objects.get(id=LessonID)
    else:
        lesson = None
    
    #Data
    #data = AcceptanceJoinTablesX2(LessonID).order_by('LessonID__name','SchoolToGradeID__name')
    data = None
    #print data    


    """
    # Prepare document for download        
    # -----------------------------
    f = StringIO()
    document.save(f)
    length = f.tell()
    f.seek(0)
    response = HttpResponse(
        f.getvalue(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    response['Content-Disposition'] = 'attachment; filename=' + docx_title
    response['Content-Length'] = length
    return response
    """
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename=docTest.docx'
    docx_data = doDocxTest(data, lesson)
    response.write(docx_data)
    return response


#########################################
# DOCX Letters
#########################################
def doDocxTest(dataDB, lesson=None):
    
    #global lesson_col_width    
    #output = StringIO.StringIO()    
    #document = Document(output)
    doc = Document()
    docx_title="TEST_DOCUMENT.docx"
            
    # ---- Cover Letter ----
    #document.add_picture((r'%s/static/images/my-header.png' % (settings.PROJECT_PATH)), width=Inches(4))
    doc.add_paragraph()
    doc.add_paragraph("%s" % date.today().strftime('%B %d, %Y'))
    doc.add_paragraph('Dear Sir or Madam:')
    doc.add_paragraph('We are pleased to help you with your widgets.')
    doc.add_paragraph('Please feel free to contact me for any additional information.')
    doc.add_paragraph('I look forward to assisting you in this project.')
    doc.add_page_break()

    # Prepare document for download        
    # -----------------------------
    output = StringIO.StringIO()
    #output = StringIO()
    doc.save(output)
    length = output.tell()
    output.seek(0)
    
    """
    #Response
    response = HttpResponse(
        output.getvalue(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    response['Content-Disposition'] = 'attachment; filename=' + docx_title
    response['Content-Length'] = length
    return response

    """

    # close workbook
    #workbook.close()
    docx_data = output.getvalue()
    return docx_data

#############################################
# ΔΙΑΒΙΒΑΣΤΙΚΑ ΣΧΟΛΕΙΩΝ για ΑΠΟΣΤΟΛΗ
#############################################
def docSchoolCoverLetter(request):
    
    SchoolToGradeId = request.GET.get('SchoolToGradeID', None)
    print SchoolToGradeId
    
    if SchoolToGradeId: 
        #theSchoolToGrade = SchoolToGrade.objects.get(id=SchoolToGradeId)    
        data = SchoolToGrade.objects.filter(id=SchoolToGradeId)
        #data = data.filter(LessonID=LessonID)
    else:
        data = SchoolToGrade.objects.all()
 
    data = data.order_by('name')
    #print data
    
    #PDF part 
    #Buffer
    #buffer = BytesIO()
    #pdf = doPDFSchoolCoverLetter(buffer, data, LetterCode=1)

    #DOC part 
    docx_data = doDocSchoolCoverLetter(data)

    #Response part
    #filename = 'Letters' + today.strftime('%Y-%m-%d')    
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    today = date.today()
    filename = 'SchoolCoverLetters'
    response['Content-Disposition'] = 'attachement; filename={0}.docx'.format(filename)
    
    response.write(docx_data)
    return response


#########################################
# DOCX Letters
#########################################
def doDocSchoolCoverLetter(dataDB):
    
    #global lesson_col_width    
    #output = StringIO.StringIO()    
    #document = Document(output)
    doc = Document()
    section = doc.sections[0]
    section.bottom_margin = Mm(10)
    section.top_margin = Mm(10)
    section.left_margin = Mm(10)
    section.right_margin = Mm(10)
    
    # Get GC info record
    gcinfo = GradeCenterInfo.getGCInfo(GradeCenterInfo)
    #record.update( name, article,  presidentName, presidentSurname, phone, folderBooks, minFolderBooks, 
    constGCName=gcinfo.name
    constGCPresdArticle=gcinfo.article
    constGCPresdName=gcinfo.presidentName
    constGCPresdSurname=gcinfo.presidentSurname
    date_text =  datetime.now().strftime('%d/%m/%Y %H:%M')
    
    #docx_title=u'ΔΙΑΒΙΒΑΣΤΙΚΟ ΣΧΟΛΕΙΩΝ.docx'    
    
    # ---- Cover Letter ----
    #document.add_picture((r'%s/static/images/my-header.png' % (settings.PROJECT_PATH)), width=Inches(4))
    #doc.add_paragraph()
    #doc.add_paragraph("%s" % date.today().strftime('%B %d, %Y'))
    #logo = settings.STATIC_ROOT + "/static/images/" + "python_logo.png"    
    formatted_time = time.ctime()


    #Rec.-based data
    row = 0 
    for school in dataDB:
        #StoryDB = []        
        name = school.name
        ddeName = school.ddeName
        #print "%s %s %s %s" %(s.code, s.name, s.ddeCode, s.ddeName)
        #ptext2 = u'<font name="DejaVuSansMono"> Σήμερα %s ημέρα  %s και ώρα %s</font>' %(datetime.now(), datetime.now(), datetime.now())

        #Titles etc
        doc.add_heading(constGCName, level=1)    
        doc.add_heading(u'ΠΡΩΤΟΚΟΛΛΟ ΠΑΡΑΔΟΣΗΣ ΚΑΙ ΠΑΡΑΛΑΒΗΣ', 0)    
        #doc.add_paragraph('Intense quote', style='IntenseQuote')        
        #List & Bullet paragaphs
        #doc.add_paragraph('first item in unordered list', style='ListBullet')
        #doc.add_paragraph('first item in ordered list', style='ListNumber')        
        #Image
        #doc.add_picture('python.jpeg', width=Inches(1.25))

        paragraph = doc.add_paragraph()
        text = u'Σήμερα ............... ημέρα ............... και ώρα .............'
        run = paragraph.add_run(text)
        text = u'%s υπογραφόμεν%s Πρόεδρος της Επιτροπής του %s παρέδωσα στη ' %(constGCPresdArticle, constGCPresdArticle, constGCName)
        run = paragraph.add_run(text)
        text = u'Δ/νση Δ.Ε. %s ' %(ddeName)
        run = paragraph.add_run(text)
        run.bold = True
        text = u'τα αποκόμματα, των γραπτών δοκιμίων και απουσιών, τις καταστάσεις βαθμολογίας '
        run = paragraph.add_run(text)
        text = u'των μαθητών της Γ΄Τάξης του Σχολείου με την Επωνυμία '
        run = paragraph.add_run(text)
        text = u'%s ' %(name)
        run = paragraph.add_run(text)
        run.bold = True
        text = u'που εξετάστηκαν στο %s.' %(constGCName)
        run = paragraph.add_run(text)
        
        """
        p = doc.add_paragraph(text)
        p.add_run('bold').bold = True
        p.add_run(' and some ')
        p.add_run('italic.').italic = True
        """
        
        text = u'Ο αριθμός τους κατά μάθημα και κατεύθυνση φαίνεται στον παρακάτω πίνακα.'
        doc.add_paragraph(text)

        #Make TABLE 
        #table_data = [(u'ΜΑΘΗΜΑ',u'ΤΕΤΡΑΔΙΑ',u'ΑΠΟΥΣΙΕΣ',u'ΜΗΔΕΝΙΣΜΕΝΑ', u'ΣΥΝΟΛΑ', )]

        # Create the table
        #-----------------
        #table = Table(table_data, rowHeights=20)
        table = doc.add_table(rows=1, cols=5)

        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = u'ΜΑΘΗΜΑ'
        hdr_cells[1].text = u'ΤΕΤΡΑΔΙΑ'
        hdr_cells[2].text = u'ΑΠΟΥΣΙΕΣ'
        hdr_cells[3].text = u'ΜΗΔΕΝΙΣΜΕΝΑ'
        hdr_cells[4].text = u'ΣΥΝΟΛΑ'
        
        #Loop RECORDs
        for rec in school.acceptance_set.all():
            # calc partial sum
            #table_data = table_data + [(rec.LessonID.name, rec.books, rec.booksAbscent, rec.booksZero, part_sum)]
            row_cells = table.add_row().cells
            row_cells[0].text = rec.LessonID.name
            row_cells[1].text = str(rec.books)
            row_cells[2].text = str(rec.booksAbscent)
            row_cells[3].text = str(rec.booksZero)            
            part_sum = (rec.books + rec.booksAbscent + rec.booksZero)
            row_cells[4].text = str(part_sum)


        # President Signature
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        text = u'%s Πρόεδρος του %s' %(constGCPresdArticle, constGCName)
        run = paragraph.add_run(text)
        run.add_break()        
        run.add_break()        
        run.add_break()        
        text = u'%s %s' %(constGCPresdName, constGCPresdSurname)
        run = paragraph.add_run(text)
        
        #Page break        
        doc.add_page_break()
    

    # Prepare document for download        
    # -----------------------------
    output = StringIO.StringIO()
    #output = StringIO()
    doc.save(output)
    length = output.tell()
    output.seek(0)
    
    """
    #Response
    response = HttpResponse(
        output.getvalue(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    response['Content-Disposition'] = 'attachment; filename=' + docx_title
    response['Content-Length'] = length
    return response

    """


    # close workbook
    #workbook.close()
    docx_data = output.getvalue()
    return docx_data


